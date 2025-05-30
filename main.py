import os
import time
import logging
import streamlit as st
from openai import AzureOpenAI
from PyPDF2 import PdfReader
import docx
import difflib
from dotenv import load_dotenv
import re
from io import BytesIO

def get_word_changes(original, improved):
    original_words = re.findall(r'\w+|\W', original)
    improved_words = re.findall(r'\w+|\W', improved)

    sm = difflib.SequenceMatcher(None, original_words, improved_words)
    changes = []

    for opcode, i1, i2, j1, j2 in sm.get_opcodes():
        if opcode == 'replace':
            changed_from = ''.join(original_words[i1:i2]).strip()
            changed_to = ''.join(improved_words[j1:j2]).strip()
            if changed_from and changed_to:
                changes.append(f'"{changed_from}" → "{changed_to}"')
        elif opcode == 'delete':
            removed = ''.join(original_words[i1:i2]).strip()
            if removed:
                changes.append(f'Removed: "{removed}"')
        elif opcode == 'insert':
            added = ''.join(improved_words[j1:j2]).strip()
            if added:
                changes.append(f'Added: "{added}"')

    return changes

load_dotenv()

api_key = os.getenv("API_KEY")
endpoint = os.getenv("ENDPOINT")

st.set_page_config(page_title="Legal Writing Assistant", layout="wide")
logging.basicConfig(level=logging.INFO)

st.markdown(
    """
    <style>
    section[data-testid="stSidebar"] {
        background-color: rgba(0, 0, 0, 0.5);  
        color: white;
    }

    .sidebar-content {
        padding: 1rem;
        font-size: 0.9rem;
        color: white;
    }
    </style>
    """,
    unsafe_allow_html=True
)

client = AzureOpenAI(api_key=api_key, azure_endpoint=endpoint, api_version="2024-02-15-preview")

LEGAL_WRITING_GUIDELINES = """
LEGAL WRITING GUIDELINES

These 15 principles are the foundation of clear, persuasive legal writing. Apply them rigorously. Use silent, structured reasoning to guide each revision, but show only the final product—not your process.

1. BE CONCISE  
Ask: Does this word or phrase add legal value? Could this be shorter without loss of meaning?  
✘ "The argument made by opposing counsel is one that fails to succeed for reasons including, inter alia, the fact that..."  
✔ "Opposing counsel’s argument fails because..."

2. USE ACTIVE VOICE  
Ask: Who is acting here? Is the subject direct and forceful?  
✘ "The penalty was called by the referee."  
✔ "The referee called the penalty."

3. REPLACE LEGALESE WITH PLAIN ENGLISH  
Ask: Can this be simplified without altering legal meaning?  
✘ "inter alia" → ✔ "among other things"  
✘ "pursuant to" → ✔ "under"  
✘ "utilize" → ✔ "use"

4. LIMIT NOMINALIZATIONS  
Ask: Are verbs being buried as nouns? Revive them.  
✘ "There was committee agreement."  
✔ "The committee agreed."

5. OMIT UNNECESSARY WORDS  
Ask: Can this be said more directly?  
✘ "At the point in time" → ✔ "Then"  
✘ "By means of" → ✔ "By"  
✘ "In light of the fact that" → ✔ "Because"

6. AVOID RUN-ON SENTENCES  
Ask: Does this sentence contain more than one core idea or exceed 20–25 words? Break it.  
✘ Overloaded sentence → ✔ One idea per sentence.

7. BREAK LONG PARAGRAPHS  
Ask: Does this paragraph cover more than one concept? Could breaking it help clarity?  
✔ Each paragraph should focus on a single theme or step in the argument.

8. ELIMINATE REDUNDANCY  
Ask: Are there repeated ideas or synonyms? Cut them.  
✘ "Null and void"  
✔ "Void"  
✘ "Cease and desist"  
✔ "Stop"

9. REMOVE WEASEL WORDS  
Ask: Does this word add vagueness or false confidence?  
✘ "Clearly," "arguably," "somewhat"  
✔ Omit or replace with precise reasoning.

10. AVOID DOUBLE NEGATIVES  
Ask: Is this clearer in a positive form?  
✘ "Not uncommon" → ✔ "Common"  
✘ "Not without merit" → ✔ "Has merit"

11. STRIP EMPTY PHRASES  
Ask: Is this introductory phrase adding anything?  
✘ "I would like to point out that..."  
✔ "Chester v. Morris was overruled."

12. CONSIDER LEGALLY RELEVANT ALTERNATIVES  
Ask: Are multiple interpretations possible? Should they be addressed?  
✔ Present plausible alternatives where relevant, not to hedge but to strengthen analysis.

13. MARSHAL RELEVANT FACTS, LAW, AND POLICY  
Ask: Have you gathered and integrated all key materials?  
✔ Use facts, legal rules, and policy rationales to support—not just decorate—your argument.

14. THINK CRITICALLY ABOUT INFORMATION  
Ask: Are your sources valid, and have you weighed their strengths and weaknesses?  
✔ Don't just cite—analyze. Distinguish, analogize, or critique as needed.

15. REACH A DEFINITIVE CONCLUSION  
Ask: What’s the clearest, most supportable legal position?  
✔ End with a firm conclusion, even if it acknowledges limitations.

—

Apply these principles with precision and discipline. The reader should see only clear, confident legal writing. Your thoughtful process should be invisible—but your mastery, unmistakable.
"""



@st.cache_resource
def create_assistant():
    return client.beta.assistants.create(
        name="Legal Writing Assistant",
        instructions="You are a legal writing coach for law firm associates.",
        tools=[],
        model="gpt-4o"
    )

def run_assistant(prompt, task_instructions):
    try:
        assistant = create_assistant()
        thread = client.beta.threads.create()
        client.beta.threads.messages.create(thread_id=thread.id, role="user", content=prompt)
        run = client.beta.threads.runs.create(
            thread_id=thread.id,
            assistant_id=assistant.id,
            instructions=task_instructions + "\n\n" + LEGAL_WRITING_GUIDELINES
        )

        while True:
            run_status = client.beta.threads.runs.retrieve(thread_id=thread.id, run_id=run.id)
            if run_status.status in ["completed", "failed", "cancelled", "expired"]:
                break
            time.sleep(1)

        if run_status.status == "completed":
            messages = client.beta.threads.messages.list(thread_id=thread.id)
            for msg in reversed(list(messages)):
                if msg.role == "assistant":
                    return msg.content[0].text.value
            return "⚠️ No assistant response found."
        else:
            return f"❌ Run failed with status: {run_status.status}"
            #TypeError: expected string or bytes-like object, got 'NoneType'
    except Exception as e:
        logging.error("Assistant error", exc_info=True)
        return f"🚨 Error: {str(e)}"

def parse_score_from_evaluation(evaluation_text):
    import re
    match = re.search(r"score\s*[:\-]?\s*(\d{1,3})", evaluation_text, re.IGNORECASE)
    if match:
        score = int(match.group(1))
        if 0 <= score <= 100:
            return score
    return None

def highlight_differences(original, improved):
    differ = difflib.Differ()
    diff = list(differ.compare(original.splitlines(), improved.splitlines()))
    highlighted = []
    for line in diff:
        if line.startswith("  "):
            highlighted.append(line[2:])
        elif line.startswith("- "):
            highlighted.append(f'<del style="background-color:#faa;">{line[2:]}</del>')
        elif line.startswith("+ "):
            highlighted.append(f'<span style="background-color: #fffb91;">{line[2:]}</span>')
    return "<br>".join(highlighted)


def extract_text_and_metadata(file):
    try:
        metadata_cleaned = {}
        page_texts = {}

        if file.type == "application/pdf":
            reader = PdfReader(file)
            raw_metadata = reader.metadata

            
            if raw_metadata:
                metadata_cleaned = {k.lstrip("/"): v for k, v in raw_metadata.items() if v}

           
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    page_texts[f"Page {i + 1}"] = text.strip()

        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(file)
            raw_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            page_char_limit = 800
            current_page = []
            current_length = 0
            page_num = 1

            for para in raw_paragraphs:
                current_page.append(para)
                current_length += len(para)

                if current_length >= page_char_limit:
                    page_texts[f"Page {page_num}"] = "\n".join(current_page)
                    page_num += 1
                    current_page = []
                    current_length = 0

           
            if current_page:
                page_texts[f"Page {page_num}"] = "\n".join(current_page)

            
            props = doc.core_properties
            if props.title:
                metadata_cleaned["Title"] = props.title
            if props.author:
                metadata_cleaned["Author"] = props.author

        else:
            st.warning("Unsupported file type.")
            return {}, {}

        return metadata_cleaned, page_texts

    except Exception as e:
        st.error(f"❌ Failed to extract content: {e}")
        return {}, {}

st.title("Legal Writing Assistant")

with st.sidebar:
    st.markdown("### 📁 Upload Document")
    uploaded_file = st.file_uploader("Upload a legal document (.docx or .pdf)", type=["pdf", "docx"])

    if uploaded_file:
        metadata, page_texts = extract_text_and_metadata(uploaded_file)

        if not page_texts:
            st.error("❌ No valid text found in the document.")
        else:
            total_pages = metadata.get("Total Pages", len(page_texts))

            st.markdown(f"- **Total Pages**: {len(page_texts)}")

            if st.button("🧮 Evaluate and Improve"):
                full_text = "\n\n".join(page_texts.values())

                # --- AGENT 1: EVALUATION ---
                with st.spinner("Evaluating your legal writing..."):
                    evaluation = run_assistant(full_text, 
                        """You are a legal writing evaluator. Assess the following text using the 15 Legal Writing Guidelines as your sole standard. 
                        Assign a final score from 0 to 100 reflecting the overall quality of legal writing.
                        Then provide concise, bullet-point feedback grouped by:
                        - Strengths
                        - Weaknesses
                        - Actionable Suggestions
                        Format:
                        Score: XX
                        - Strength: ...
                        - Weakness: ...
                        - Suggestion: ...
                        """
                    )

                score = parse_score_from_evaluation(evaluation)

                if score is None:
                    st.warning("⚠️ Could not detect score.")
                elif score >= 95:
                    st.success(f"✅ Score: {score}/100 — Your writing is strong. No improvement needed.")
                else:
                    st.warning(f"⚠️ Score: {score}/100 — Improvements recommended.")
                    st.session_state["show_improvements"] = True
                    st.session_state["page_texts"] = page_texts

if st.session_state.get("show_improvements"):
    page_texts = st.session_state["page_texts"]
    tabs = st.tabs([f"{page_num}" for page_num in page_texts.keys()])
    page_summaries = []
    
    for idx, page_num in enumerate(page_texts.keys()):
        original = page_texts[page_num]
        if not original.strip():
            continue
        
        improve_instructions = (
            "You are a meticulous legal writing assistant. Silently apply the 15 Legal Writing Guidelines to evaluate and improve the following text. "
            "Use internal reasoning, but do not display your thought process.\n\n"
            "Focus on:\n"
            "• Concise, direct language (remove redundancies, simplify phrases, eliminate legalese)\n"
            "• Active voice and clear verb use\n"
            "• Shorter, well-structured sentences and logically focused paragraphs\n"
            "• A professional, precise, and confident tone throughout\n"
            "• Legal integrity—preserve the original legal meaning at all times\n\n"
            "Output only:\n"
            "The final, revised version of the text—polished and professional\n"
        )
        
        format_instructions = ( " Clean the text by removing headers, footers, and page numbers."
                                "Organize content using headings, short paragraphs, and bullet points."
                                "if breakpoint put it in new line"
                                )
        format_original = run_assistant(original, format_instructions)
        
        # --- AGENT 2: IMPROVE ---
        improved = run_assistant(original, improve_instructions)
        
        with tabs[idx]:
            st.markdown(f"### 🧾 {page_num} Comparison")
            col1, col2 = st.columns(2)

            with col1:
                st.subheader("📄 Original Text")
                st.text_area(label="", value=format_original.strip(), height=2000, key=f"original_{page_num}")

            with col2:
                st.subheader("✅ Improved Text")
                if f"improved_{page_num}" not in st.session_state:
                    st.session_state[f"improved_{page_num}"] = improved.strip()
                st.text_area(label="", value=improved.strip(), height=2000, key=f"improved_{page_num}")

            with st.expander("🔍 Word-Level Changes"):
                word_changes = get_word_changes(original, improved)
                if word_changes:
                    summary = f"Page {page_num}:\n" + "\n".join(f"- {change}" for change in word_changes)
                    st.markdown(summary)
                else:
                    summary = f"Page {page_num}: No changes detected."
                    st.markdown(summary)

            page_summaries.append(summary)
    st.session_state["all_pages_processed"] = True

def create_docx_from_improved_text():
    doc = docx.Document()
    doc.add_heading("Improved Legal Document", level=1)

    for page_num in st.session_state.get('page_texts', {}).keys():
        improved_text = st.session_state.get(f"improved_{page_num}", "")
        if improved_text.strip():
            lines = improved_text.split('\n')

            buffer = []
            for line in lines:
                if line.strip():
                    buffer.append(line.strip())
                elif buffer:
                    doc.add_paragraph('\n'.join(buffer))
                    buffer = []
            if buffer:
                doc.add_paragraph('\n'.join(buffer))

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

if st.session_state.get('all_pages_processed', False):
    docx_io = create_docx_from_improved_text()
    st.sidebar.download_button(
        label="📄 Download Improved Document `.docx`",
        data=docx_io,
        file_name="improved_legal_document.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
