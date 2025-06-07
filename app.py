# --- IMPORTS ---

import os
import re
import time
import logging
import difflib
import docx
from io import BytesIO
from PyPDF2 import PdfReader
import streamlit as st
from dotenv import load_dotenv
from openai import AzureOpenAI


# --- UTILITIES ---

def get_word_changes(original: str, improved: str) -> list[str]:
    """
    Compute word-level changes between original and improved texts.
    Returns a list of human-readable change descriptions.
    """
    original_words = re.findall(r"\w+|\W", original)
    improved_words = re.findall(r"\w+|\W", improved)

    sm = difflib.SequenceMatcher(None, original_words, improved_words)
    changes = []

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "replace":
            changes.append(
                f'"{"".join(original_words[i1:i2])}" → '
                f'"{"".join(improved_words[j1:j2])}"'
            )
        elif tag == "delete":
            changes.append(f'Removed: "{"".join(original_words[i1:i2])}"')
        elif tag == "insert":
            changes.append(f'Added: "{"".join(improved_words[j1:j2])}"')

    return changes


def parse_score_from_evaluation(evaluation_text: str) -> int | None:
    """
    Extract a numeric score (0-100) from the evaluation text.
    Returns None if no valid score found.
    """
    match = re.search(r"score\s*[:\-]?\s*(\d{1,3})", evaluation_text, re.IGNORECASE)
    if match:
        score = int(match.group(1))
        if 0 <= score <= 100:
            return score
    return None


# --- ENVIRONMENT SETUP ---
load_dotenv()
API_KEY = os.getenv("API_KEY")
ENDPOINT = os.getenv("ENDPOINT")


# --- STREAMLIT SETUP ---

st.set_page_config(page_title="Legal Writing Assistant", layout="wide")
logging.basicConfig(level=logging.INFO)


# --- AZURE OPENAI CLIENT SETUP ---

client = AzureOpenAI(
    api_key=API_KEY,
    azure_endpoint=ENDPOINT,
    api_version="2024-02-15-preview",
)


@st.cache_resource
def create_assistant():
    try:
        return client.beta.assistants.create(
            name="Legal Writing Assistant",
            instructions="You are a legal writing coach for law firm associates.",
            tools=[],
            model="gpt-4o",
        )
    except Exception as e:
        logging.error("Error creating assistant", exc_info=True)
        st.error("❌ Failed to create assistant.")
        raise

def run_assistant(prompt: str, task_instructions: str, max_retries: int = 3) -> str:
    """
    Run the Azure OpenAI assistant with retries for rate limiting.
    Returns the assistant's response or an error message.
    """
    retries = 0
    while retries < max_retries:
        try:
            assistant = create_assistant()
            thread = client.beta.threads.create()
            client.beta.threads.messages.create(
                thread_id=thread.id, role="user", content=prompt
            )

            run = client.beta.threads.runs.create(
                thread_id=thread.id,
                assistant_id=assistant.id,
                instructions=task_instructions + "\n\n" + LEGAL_WRITING_GUIDELINES,
            )

            while True:
                run_status = client.beta.threads.runs.retrieve(
                    thread_id=thread.id, run_id=run.id
                )
                if run_status.status in ["completed", "failed", "cancelled", "expired"]:
                    break
                time.sleep(1)

            if run_status.status == "completed":
                messages = client.beta.threads.messages.list(thread_id=thread.id)
                for msg in reversed(list(messages)):
                    if msg.role == "assistant":
                        try:
                            return msg.content[0].text.value
                        except (AttributeError, IndexError, TypeError) as ex:
                            logging.warning(
                                f"Failed to extract assistant response: {ex}"
                            )
                            continue
                return "⚠️ No assistant response found."

            if run_status.status == "failed":
                last_error = getattr(run_status, "last_error", None)
                if last_error:
                    if last_error.code == "rate_limit_exceeded":
                        wait_time = 10
                        logging.warning(
                            f"Rate limit exceeded. Waiting {wait_time} seconds..."
                        )
                        time.sleep(wait_time)
                        retries += 1
                        continue
                    return f"❌ Run failed: {last_error.code} - {last_error.message}"
                return "❌ Run failed with unknown error."

            return f"❌ Run status: {run_status.status}"

        except Exception as e:
            logging.exception("🚨 Exception during assistant execution.")
            return f"🚨 Exception occurred: {e}"

    return "❌ Max retries reached due to rate limits. Please try again later."


# --- FILE PARSING LOGIC ---
def extract_text_and_metadata(file) -> tuple[dict, dict]:
    metadata_cleaned = {}
    page_texts = {}

    try:
        # Try to detect file type
        file_type = getattr(file, "type", None)
        if not file_type:
            filename = getattr(file, "name", "").lower()
            if filename.endswith(".pdf"):
                file_type = "application/pdf"
            elif filename.endswith(".docx"):
                file_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        if file_type == "application/pdf":
            reader = PdfReader(file)
            raw_metadata = reader.metadata or {}
            metadata_cleaned = {k.lstrip("/"): v for k, v in raw_metadata.items() if v}

            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        page_texts[f"Page {i + 1}"] = text.strip()
                except Exception as e:
                    logging.warning(f"Failed to read Page {i+1}: {e}")

        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(file)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            page_char_limit = 800
            current_page, current_length, page_num = [], 0, 1

            for para in paragraphs:
                current_page.append(para)
                current_length += len(para)
                if current_length >= page_char_limit:
                    page_texts[f"Page {page_num}"] = "\n".join(current_page)
                    current_page, current_length, page_num = [], 0, page_num + 1

            if current_page:
                page_texts[f"Page {page_num}"] = "\n".join(current_page)

            props = doc.core_properties
            metadata_cleaned["Title"] = props.title or "Untitled"
            metadata_cleaned["Author"] = props.author or "Unknown"

        else:
            try:
                import streamlit as st
                st.warning("Unsupported file type.")
            except ImportError:
                logging.warning("Unsupported file type.")
            return {}, {}

    except Exception as e:
        try:
            import streamlit as st
            st.error(f"❌ Failed to extract content: {e}")
        except ImportError:
            logging.error(f"❌ Failed to extract content: {e}")
        logging.error("File parsing error", exc_info=True)

    return metadata_cleaned, page_texts

# --- LEGAL WRITING GUIDELINES ---

LEGAL_WRITING_GUIDELINES = """
LEGAL WRITING GUIDELINES

These 15 principles are the foundation of clear, persuasive legal writing. Apply them rigorously. Use silent, structured reasoning to guide each revision, but show only the final product—not your process.

1. BE CONCISE  
Ask: Does this word or phrase add legal value? Could this be shorter without loss of meaning?  
✘ "The argument made by opposing counsel is one that fails to succeed for reasons including, inter alia, the fact that..."  
✔ "Opposing counsel’s argument fails because..."

2. USE ACTIVE VOICE  
Ask: Who is acting here? Is the subject direct and forceful?  
✘ "The penalty was called by the referee."  
✔ "The referee called the penalty."

3. REPLACE LEGALESE WITH PLAIN ENGLISH  
Ask: Can this be simplified without altering legal meaning?  
✘ "inter alia" → ✔ "among other things"  
✘ "pursuant to" → ✔ "under"  
✘ "utilize" → ✔ "use"

4. LIMIT NOMINALIZATIONS  
Ask: Are verbs being buried as nouns? Revive them.  
✘ "There was committee agreement."  
✔ "The committee agreed."

5. OMIT UNNECESSARY WORDS  
Ask: Can this be said more directly?  
✘ "At the point in time" → ✔ "Then"  
✘ "By means of" → ✔ "By"  
✘ "In light of the fact that" → ✔ "Because"

6. AVOID RUN-ON SENTENCES  
Ask: Does this sentence contain more than one core idea or exceed 20–25 words? Break it.  
✘ Overloaded sentence → ✔ One idea per sentence.

7. BREAK LONG PARAGRAPHS  
Ask: Does this paragraph cover more than one concept? Could breaking it help clarity?  
✔ Each paragraph should focus on a single theme or step in the argument.

8. ELIMINATE REDUNDANCY  
Ask: Are there repeated ideas or synonyms? Cut them.  
✘ "Null and void"  
✔ "Void"  
✘ "Cease and desist"  
✔ "Stop"

9. REMOVE WEASEL WORDS  
Ask: Does this word add vagueness or false confidence?  
✘ "Clearly," "arguably," "somewhat"  
✔ Omit or replace with precise reasoning.

10. AVOID DOUBLE NEGATIVES  
Ask: Is this clearer in a positive form?  
✘ "Not uncommon" → ✔ "Common"  
✘ "Not without merit" → ✔ "Has merit"

11. STRIP EMPTY PHRASES  
Ask: Is this introductory phrase adding anything?  
✘ "I would like to point out that..."  
✔ "Chester v. Morris was overruled."

12. CONSIDER LEGALLY RELEVANT ALTERNATIVES  
Ask: Are multiple interpretations possible? Should they be addressed?  
✔ Present plausible alternatives where relevant, not to hedge but to strengthen analysis.

13. MARSHAL RELEVANT FACTS, LAW, AND POLICY  
Ask: Have you gathered and integrated all key materials?  
✔ Use facts, legal rules, and policy rationales to support—not just decorate—your argument.

14. THINK CRITICALLY ABOUT INFORMATION  
Ask: Are your sources valid, and have you weighed their strengths and weaknesses?  
✔ Don't just cite—analyze. Distinguish, analogize, or critique as needed.

15. REACH A DEFINITIVE CONCLUSION  
Ask: What’s the clearest, most supportable legal position?  
✔ End with a firm conclusion, even if it acknowledges limitations.

—

Apply these principles with precision and discipline. The reader should see only clear, confident legal writing. Your thoughtful process should be invisible—but your mastery, unmistakable.
"""


# --- USER INTERFACE SETUP ---
st.title("Legal Writing Assistant")

# Initialize session state if not already
if "page_texts" not in st.session_state:
    st.session_state.page_texts = {}
if "input_source" not in st.session_state:
    st.session_state.input_source = None

# Sidebar uploader
st.sidebar.markdown("### 📁 Upload Document")
uploaded_file = st.sidebar.file_uploader(
    "Upload a legal document (.docx or .pdf)",
    type=["pdf", "docx"],
    label_visibility="collapsed",
)

# Handle uploaded document
if uploaded_file:
    metadata, page_texts = extract_text_and_metadata(uploaded_file)
    st.session_state.page_texts = page_texts
    st.session_state.input_source = "document"

# Chat input at bottom
user_input = st.chat_input("Paste your legal text here", key="chat_input")

if user_input and user_input.strip():
    st.session_state.page_texts = {"Page 1": user_input.strip()}
    st.session_state.input_source = "text"

# Use the stored version
page_texts = st.session_state.page_texts
input_source = st.session_state.input_source

# Display info
if page_texts:
    total_pages = len(page_texts)
    st.sidebar.markdown(f"- **Total Pages**: {total_pages}")

    if st.sidebar.button("🧮 Evaluate and Improve"):
        full_text = "\n\n".join(page_texts.values())

        with st.spinner("Evaluating your legal writing..."):
            evaluation = run_assistant(
                full_text,
                (
                    "You are a legal writing evaluator. Your task is to evaluate ONLY legal writing. "
                    "If the input does not appear to be legal writing (e.g., casual text, fiction, code, general information), "
                    "respond with: 'The input does not appear to be legal writing and will not be evaluated.'\n\n"
                    "If the input is legal in nature, assess it strictly based on the 15 Legal Writing Guidelines. "
                    "Assign a final score from 0 to 100 reflecting the overall quality of legal writing.\n\n"
                    "Then provide concise, bullet-point feedback grouped by:\n"
                    "- Strengths\n"
                    "- Weaknesses\n"
                    "- Actionable Suggestions\n\n"
                    "Format:\n"
                    "Score: XX\n"
                    "- Strength: ...\n"
                    "- Weakness: ...\n"
                    "- Suggestion: ..."
                ),
            )

        score = parse_score_from_evaluation(evaluation)

        if score is None:
            st.warning("⚠️ Could not detect score.")
        elif score >= 95:
            st.success(
                f"✅ Score: {score}/100 — Your writing is strong. No improvement needed."
            )
        else:
            st.warning(f"⚠️ Score: {score}/100 — Improvements recommended.")
            st.session_state["show_improvements"] = True


# --- TEXT IMPROVEMENT SECTION ---

if st.session_state.get("show_improvements"):
    page_texts = st.session_state["page_texts"]
    tabs = st.tabs([f"{page}" for page in page_texts.keys()])

    for idx, page_num in enumerate(page_texts.keys()):
        original = page_texts[page_num]
        if not original.strip():
            continue

        improve_instructions = (
            "You are a meticulous legal writing assistant. Silently apply the 15 Legal Writing Guidelines to evaluate and improve the following text. "
            "Use internal reasoning, but do not display your thought process.\n\n"
            "Focus on:\n"
            "• Concise, direct language (remove redundancies, simplify phrases, eliminate legalese)\n"
            "• Active voice and clear verb use\n"
            "• Shorter, well-structured sentences and logically focused paragraphs\n"
            "• A professional, precise, and confident tone throughout\n"
            "• Legal integrity—preserve the original legal meaning at all times\n\n"
            "Output only:\n"
            "The final, revised version of the text—polished and professional\n"
        )

        format_instructions = (
            "Clean the text by removing headers, footers, and page numbers. "
            "Organize content using headings, short paragraphs, and bullet points. "
            "If breakpoint put it in new line."
        )

        format_original = run_assistant(original, format_instructions)
        improved = run_assistant(original, improve_instructions)

        with tabs[idx]:
            st.markdown(f"### 🧾 {page_num} Comparison")
            col1, col2 = st.columns(2)

            with col1:
                st.subheader("📄 Original Text")
                cleaned_original = re.sub(r"(\*\*|__|##+)", "", format_original)
                st.text_area(
                    "",
                    value=cleaned_original.strip(),
                    height=2000,
                    key=f"original_{page_num}",
                )

            with col2:
                st.subheader("✅ Improved Text")
                st.text_area(
                    "", value=improved.strip(), height=2000, key=f"improved_{page_num}"
                )
                if f"improved_{page_num}" not in st.session_state:
                    st.session_state[f"improved_{page_num}"] = improved.strip()

            with st.expander("🔍 Word-Level Changes"):
                word_changes = get_word_changes(original, improved)
                if word_changes:
                    summary = f"Page {page_num}:\n" + "\n".join(
                        f"- {c}" for c in word_changes
                    )
                    st.markdown(summary)
                else:
                    st.markdown(f"Page {page_num}: No changes detected.")


# --- DOCX EXPORT SECTION ---

def create_docx_from_improved_text() -> BytesIO:
    try:
        doc = docx.Document()
        doc.add_heading("Improved Legal Document", level=1)

        for page_num in st.session_state.get("page_texts", {}):
            improved_text = st.session_state.get(f"improved_{page_num}", "")
            if improved_text.strip():
                lines = improved_text.split("\n")
                buffer = []
                for line in lines:
                    if line.strip():
                        buffer.append(line.strip())
                    elif buffer:
                        doc.add_paragraph("\n".join(buffer))
                        buffer = []
                if buffer:
                    doc.add_paragraph("\n".join(buffer))

        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io

    except Exception as e:
        logging.error("Failed to create DOCX", exc_info=True)
        st.error("❌ Could not generate DOCX file.")
        return BytesIO()


if st.session_state.get("show_improvements") and st.session_state.get("page_texts"):
    docx_io = create_docx_from_improved_text()
    st.sidebar.download_button(
        label="📄 Download Improved Document `.docx`",
        data=docx_io,
        file_name="improved_legal_document.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
